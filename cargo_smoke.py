import json, os
from pathlib import Path
from docx import Document

tmp = Path("dist") / "_exe_smoke_policy"
tmp.mkdir(exist_ok=True)
for name, content in [("Ген. полис.docx", "Условия полиса."), ("ДС №1 от 01.01.2025.docx", "ДС условия.")]:
    d = Document(); d.add_paragraph(content); d.save(tmp / name)
decl = tmp / "200.docx"
d = Document(); d.add_paragraph("Декларация: объект страхования — оборудование."); d.save(decl)
payload = {"request_id": 1, "policy_folder": str(tmp.resolve()), "declaration_paths": [str(decl.resolve())]}
Path("cargo_payload.json").write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
