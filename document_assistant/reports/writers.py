import shutil
from abc import ABC, abstractmethod
from pathlib import Path

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from document_assistant.reports.report_models import InsuranceReport, ReportRow


# ── Colour palette ────────────────────────────────────────────────────────────

_GREEN  = "D6F0D6"   # Есть
_RED    = "F0D6D6"   # Нет
_YELLOW = "FFF3CD"   # Частично
_HEADER = "1F4E79"   # dark blue for header background

_STATUS_FILL = {
    "есть":      _GREEN,
    "нет":       _RED,
    "частично":  _YELLOW,
}


def _status_fill(status: str) -> str:
    return _STATUS_FILL.get(status.lower().strip(), "FFFFFF")


# ── Abstract base ─────────────────────────────────────────────────────────────

class ReportWriter(ABC):
    """Write an InsuranceReport to a file and return the output path."""

    HEADERS = [
        "Требование клиента",
        "Покрытие по программе",
        "Статус",
        "Комментарий",
    ]

    @abstractmethod
    def write(self, report: InsuranceReport, output_path: Path, source_path: Path = None) -> Path:
        pass


# ── Excel ─────────────────────────────────────────────────────────────────────

class ExcelReportWriter(ReportWriter):

    def write(self, report: InsuranceReport, output_path: Path, source_path: Path = None) -> Path:
        if source_path and source_path.suffix.lower() in (".xlsx", ".xls"):
            return self._write_annotated(report, output_path, source_path)
        return self._write_new(report, output_path)

    def _write_new(self, report: InsuranceReport, output_path: Path) -> Path:
        wb = Workbook()
        ws = wb.active
        ws.title = "Сравнение"

        self._write_header_row(ws)
        self._write_data_rows(ws, report.rows)
        self._apply_column_widths(ws)

        if report.summary:
            self._write_summary_sheet(wb, report.summary)

        wb.save(output_path)
        return output_path

    def _write_annotated(self, report: InsuranceReport, output_path: Path, source_path: Path) -> Path:
        """Copy source file and write 3 annotation columns to all sheets.

        Rows are matched by text content across all worksheets so that:
        - rows skipped/merged by the LLM do not shift subsequent annotations
        - multi-sheet Excel files are handled correctly
        """
        shutil.copy2(source_path, output_path)
        wb = load_workbook(output_path)

        ann_col = 2
        new_headers = ["Покрытие по программе", "Статус", "Комментарий"]

        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill("solid", fgColor=_HEADER)
        center = Alignment(horizontal="center", vertical="center", wrap_text=True)
        wrap = Alignment(vertical="top", wrap_text=True)
        thin = Side(style="thin", color="CCCCCC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        # Build global index across ALL sheets: norm_text → (worksheet, row_idx)
        global_index: dict[str, tuple] = {}
        total_source_rows = 0
        for ws in wb.worksheets:
            for row_idx in range(2, ws.max_row + 1):
                val = ws.cell(row=row_idx, column=1).value
                if val:
                    key = self._norm(str(val))
                    if key not in global_index:  # first sheet wins on collision
                        global_index[key] = (ws, row_idx)
                        total_source_rows += 1

        # Write each LLM response to its matched sheet+row
        # used_locations tracks already-annotated rows to prevent double-writes
        matched = 0
        used_locations: set = set()
        sheets_touched: set = set()
        for row in report.rows:
            result = self._find_row_global(row.client_requirement, global_index)
            if result is None:
                continue
            ws, row_idx = result
            loc_key = (id(ws), row_idx)
            if loc_key in used_locations:
                continue
            used_locations.add(loc_key)
            matched += 1
            sheets_touched.add(ws)

            cov_cell = ws.cell(row=row_idx, column=ann_col, value=row.program_coverage)
            cov_cell.alignment = wrap
            cov_cell.border = border

            status_cell = ws.cell(row=row_idx, column=ann_col + 1, value=row.status)
            status_cell.fill = PatternFill("solid", fgColor=_status_fill(row.status))
            status_cell.alignment = wrap
            status_cell.border = border

            comment_cell = ws.cell(row=row_idx, column=ann_col + 2, value=row.comment)
            comment_cell.alignment = wrap
            comment_cell.border = border

        # Add annotation headers to every sheet that received annotations
        for ws in sheets_touched:
            for i, title in enumerate(new_headers):
                col = ann_col + i
                cell = ws.cell(row=1, column=col, value=title)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center
                ws.column_dimensions[cell.column_letter].width = [45, 14, 50][i]

        print(
            f"[INFO] Аннотировано {matched}/{len(report.rows)} строк LLM "
            f"из {total_source_rows} строк оригинала "
            f"({len(sheets_touched)} лист(ов))",
            flush=True,
        )

        if report.summary:
            self._write_summary_sheet(wb, report.summary)

        wb.save(output_path)
        return output_path

    @staticmethod
    def _norm(text: str) -> str:
        """Normalise requirement text for matching: lowercase, collapse spaces."""
        return " ".join(text.lower().split())

    @staticmethod
    def _words(text: str) -> set[str]:
        import re as _re
        return set(w for w in _re.split(r"\W+", text.lower()) if len(w) > 2)

    def _find_row_global(self, requirement: str, index: dict[str, tuple]) -> tuple | None:
        """Find (worksheet, row_idx) by exact, suffix, prefix, or word-overlap match.

        Matching levels (first hit wins):
        1. Exact — normalized strings are equal.
        2. Suffix — source key is the trailing part of the LLM key, separated by
           a punctuation char.  Handles the common case where the LLM prepends
           a section-header context to the actual sub-item text, e.g.
           LLM: "Первичный…приёмы: аллерголог-иммунолог"
           Source: "аллерголог-иммунолог"
        3. Prefix — one string is a prefix of the other (LLM truncation).
        4. Word-overlap ≥ 75 % — only for pairs whose word-count ratio is ≤ 3:1
           to avoid false positives from long shared prefixes.
        """
        key = self._norm(requirement)
        if not key or len(key) < 5:
            return None

        # 1. Exact match
        if key in index:
            return index[key]

        # 2. Suffix match: source row is a trailing portion of the LLM key
        #    Threshold ≥ 6 catches short specialist names (e.g. "невролог" = 8 chars)
        for orig_key, location in index.items():
            if len(orig_key) >= 6 and len(orig_key) < len(key):
                if key.endswith(orig_key):
                    sep_idx = len(key) - len(orig_key) - 1
                    if sep_idx < 0 or key[sep_idx] in (' ', ':', ',', ';', '.'):
                        return location

        # 3. Prefix match — only the direction where LLM truncated the source row
        #    (orig starts with key).  The reverse direction (key starts with orig)
        #    is excluded: that case means the LLM prepended section context, which
        #    is already handled by the suffix match above.
        for orig_key, location in index.items():
            if orig_key.startswith(key) and len(key) >= 10:
                return location

        # 4. Word-overlap ≥ 75 % — skip pairs where one text has 3× more words
        #    than the other; those share a common prefix and would cause false matches
        key_words = self._words(key)
        if len(key_words) < 3:
            return None
        best_score, best_loc = 0.0, None
        for orig_key, location in index.items():
            orig_words = self._words(orig_key)
            if not orig_words:
                continue
            shorter = min(len(key_words), len(orig_words))
            longer  = max(len(key_words), len(orig_words))
            if longer > 3 * shorter:   # extreme length mismatch → skip
                continue
            overlap = len(key_words & orig_words) / shorter
            if overlap > best_score:
                best_score, best_loc = overlap, location
        if best_score >= 0.75:
            return best_loc
        return None

    def _write_header_row(self, ws) -> None:
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill("solid", fgColor=_HEADER)
        center = Alignment(horizontal="center", vertical="center", wrap_text=True)

        for col, title in enumerate(self.HEADERS, start=1):
            cell = ws.cell(row=1, column=col, value=title)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center

    def _write_data_rows(self, ws, rows: list[ReportRow]) -> None:
        wrap = Alignment(vertical="top", wrap_text=True)
        thin = Side(style="thin", color="CCCCCC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        for row_idx, row in enumerate(rows, start=2):
            values = [
                row.client_requirement,
                row.program_coverage,
                row.status,
                row.comment,
            ]
            fill_color = _status_fill(row.status)
            fill = PatternFill("solid", fgColor=fill_color)

            for col, value in enumerate(values, start=1):
                cell = ws.cell(row=row_idx, column=col, value=value)
                cell.alignment = wrap
                cell.border = border
                if col == 3:          # Status column — coloured
                    cell.fill = fill

    @staticmethod
    def _apply_column_widths(ws) -> None:
        widths = [45, 45, 14, 50]
        for col, width in enumerate(widths, start=1):
            ws.column_dimensions[
                ws.cell(row=1, column=col).column_letter
            ].width = width

    @staticmethod
    def _write_summary_sheet(wb: Workbook, summary: str) -> None:
        ws = wb.create_sheet(title="Резюме")
        ws.column_dimensions["A"].width = 100
        cell = ws.cell(row=1, column=1, value=summary)
        cell.alignment = Alignment(wrap_text=True, vertical="top")
        ws.row_dimensions[1].height = max(15 * summary.count("\n") + 15, 40)


# ── Word ──────────────────────────────────────────────────────────────────────

class WordReportWriter(ReportWriter):

    def write(self, report: InsuranceReport, output_path: Path, source_path: Path = None) -> Path:
        doc = Document()

        heading = doc.add_heading("Анализ страхового покрытия", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._write_table(doc, report.rows)

        if report.summary:
            doc.add_heading("Резюме", level=2)
            doc.add_paragraph(report.summary)

        doc.save(output_path)
        return output_path

    def _write_table(self, doc: Document, rows: list[ReportRow]) -> None:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        for i, title in enumerate(self.HEADERS):
            hdr_cells[i].text = title
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            self._set_cell_background(hdr_cells[i], _HEADER)

        # Data rows
        _RGB = {
            _GREEN:  (0xD6, 0xF0, 0xD6),
            _RED:    (0xF0, 0xD6, 0xD6),
            _YELLOW: (0xFF, 0xF3, 0xCD),
        }

        for row in rows:
            cells = table.add_row().cells
            cells[0].text = row.client_requirement
            cells[1].text = row.program_coverage
            cells[2].text = row.status
            cells[3].text = row.comment

            fill = _status_fill(row.status)
            self._set_cell_background(cells[2], fill)

            for cell in cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

    @staticmethod
    def _set_cell_background(cell, hex_color: str) -> None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)
