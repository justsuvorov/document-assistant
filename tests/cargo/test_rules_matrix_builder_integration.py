"""End-to-end check that RulesMatrixBuilder's per-document extraction really
goes through AIAssistantService (retry/merge included) and that the
"latest ДС wins" precedence still comes out correctly once the LLM call is
routed through the shared orchestrator instead of a bespoke loop.
"""
from pathlib import Path

import pytest
from docx import Document

from document_assistant.ai.model import AIModel
from document_assistant.cargo.models import PolicySource
from document_assistant.cargo.policy_discovery import PolicyFolderScanner
from document_assistant.cargo.rules_matrix_builder import RulesMatrixBuilder


class StubMatrixModel(AIModel):
    """Returns a clause table whose text depends on call order, so the test
    can assert the LATEST ДС's extraction is what survives the merge."""

    def __init__(self):
        self.calls = 0

    def response(self, query: str) -> str:
        self.calls += 1
        text = {1: "оборудование", 2: "оборудование и запчасти", 3: "оборудование, запчасти и материалы"}[self.calls]
        return (
            "| Пункт (номер/название) | Актуальный текст/значение | Комментарий |\n"
            "|---|---|---|\n"
            f"| 9. Объект страхования | {text} | |\n"
        )


def _make_policy_layout(folder: Path) -> None:
    d = Document()
    d.add_paragraph("Условия генерального полиса.")
    d.save(folder / "ГП страхования грузов.docx")

    ds_dir = folder / "ДС"
    ds_dir.mkdir()
    for name, content in [
        ("ДС 1 (п.9).docx", "Дополнительное соглашение 1."),
        ("ДС 2 (п.9).docx", "Дополнительное соглашение 2."),
    ]:
        d = Document()
        d.add_paragraph(content)
        d.save(ds_dir / name)


class TestRulesMatrixBuilderThroughAIAssistantService:
    def test_latest_ds_wins_end_to_end(self, tmp_path: Path):
        _make_policy_layout(tmp_path)
        sources = PolicyFolderScanner().scan(str(tmp_path))
        assert [s.kind for s in sources] == ["policy", "ds", "ds"]

        model = StubMatrixModel()
        builder = RulesMatrixBuilder(model=model)

        matrix = builder.build(str(tmp_path), sources)

        assert len(matrix.clauses) == 1
        assert matrix.clauses[0].effective_text == "оборудование, запчасти и материалы"
        assert "2" in matrix.clauses[0].source_label
        assert model.calls == 3  # one call per source document

    def test_debug_files_written_next_to_each_source(self, tmp_path: Path):
        """AIAssistantService's debug/JSON cache writers should fire for each
        policy/ДС document, same as they do for the DMS pipeline's client file."""
        _make_policy_layout(tmp_path)
        sources = PolicyFolderScanner().scan(str(tmp_path))

        RulesMatrixBuilder(model=StubMatrixModel()).build(str(tmp_path), sources)

        debug_files = list(tmp_path.rglob("*_llm_debug.md"))
        json_files = list(tmp_path.rglob("*_llm_output.json"))
        assert len(debug_files) == 3
        assert len(json_files) == 3


class TestUnreadableSourceResilience:
    """A single unreadable source (a stray Office lock file, a corrupted
    upload) must not lose the whole reconciliation run."""

    def test_broken_source_is_skipped_and_run_continues(self, tmp_path: Path):
        d = Document()
        d.add_paragraph("Условия генерального полиса.")
        d.save(tmp_path / "ГП полис.docx")
        (tmp_path / "broken.docx").write_text("not a real docx", encoding="utf-8")

        sources = [
            PolicySource(kind="policy", file_path=str(tmp_path / "ГП полис.docx")),
            PolicySource(kind="ds", file_path=str(tmp_path / "broken.docx"), ds_number=1),
        ]

        matrix = RulesMatrixBuilder(model=StubMatrixModel()).build(str(tmp_path), sources)

        assert len(matrix.clauses) == 1   # policy still processed

    def test_raises_only_when_every_source_fails(self, tmp_path: Path):
        (tmp_path / "broken.docx").write_text("not a real docx", encoding="utf-8")
        sources = [PolicySource(kind="policy", file_path=str(tmp_path / "broken.docx"))]

        with pytest.raises(RuntimeError, match="Не удалось обработать ни один документ"):
            RulesMatrixBuilder(model=StubMatrixModel()).build(str(tmp_path), sources)
