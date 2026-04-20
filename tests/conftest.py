"""Shared pytest fixtures: create minimal mock files at test time."""
import os
import pytest
from pathlib import Path

# Set required env vars before any app module is imported (settings validation)
os.environ.setdefault("NORMATIVE_BASE", "/tmp/test_normative")
os.environ.setdefault("AI_ROLE", "Test role")
os.environ.setdefault("AI_PROMPT_TEMPLATE", "{role} {normative_base} {examples} {source_text}")

from openpyxl import Workbook
from docx import Document


@pytest.fixture()
def excel_file(tmp_path: Path) -> Path:
    """Excel file with one non-empty sheet and one empty sheet."""
    path = tmp_path / "test.xlsx"
    wb = Workbook()

    ws = wb.active
    ws.title = "Данные"
    ws.append(["Страховой случай", "Покрытие", "Лимит"])
    ws.append(["ДТП", "Полное", "1 000 000"])
    ws.append(["Пожар", "Частичное", "500 000"])

    ws_empty = wb.create_sheet("Пустой лист")
    # intentionally leave it empty

    wb.save(path)
    return path


@pytest.fixture()
def excel_file_all_empty(tmp_path: Path) -> Path:
    """Excel file where every sheet is empty."""
    path = tmp_path / "empty.xlsx"
    wb = Workbook()
    wb.active.title = "Пустой"
    wb.save(path)
    return path


@pytest.fixture()
def word_file(tmp_path: Path) -> Path:
    """Word file with a heading, a paragraph, and a table."""
    path = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Требования к страхованию", level=1)
    doc.add_paragraph("Клиент запрашивает страхование имущества.")

    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Риск"
    table.cell(0, 1).text = "Лимит"
    table.cell(1, 0).text = "Пожар"
    table.cell(1, 1).text = "2 000 000"
    table.cell(2, 0).text = "Затопление"
    table.cell(2, 1).text = "1 000 000"

    doc.save(path)
    return path


@pytest.fixture()
def normative_txt(tmp_path: Path) -> Path:
    """Plain-text normative base file."""
    path = tmp_path / "normative.txt"
    path.write_text("# Программа А\nПокрывает пожар и ДТП.\n", encoding="utf-8")
    return path


@pytest.fixture()
def examples_dir(tmp_path: Path, excel_file: Path, word_file: Path) -> Path:
    """Examples directory with one valid example sub-folder."""
    ex_dir = tmp_path / "examples" / "001"
    ex_dir.mkdir(parents=True)
    # copy fixtures into the sub-folder
    import shutil
    shutil.copy(excel_file, ex_dir / "client.xlsx")
    shutil.copy(word_file,  ex_dir / "response.docx")
    return tmp_path / "examples"
