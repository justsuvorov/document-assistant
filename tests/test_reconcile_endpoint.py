"""End-to-end test of POST /api/reconcile through the real FastAPI app —
not just the internal cargo classes — to lock down that the endpoint truly
builds and runs AIAssistantService per declaration (mirroring _build_service
for /api/update), with no separate orchestrator class in between.
"""
import re
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

import document_assistant.ai.model as model_mod
from document_assistant.ai.model import AIModel


class StubModel(AIModel):
    """Returns a clause table for matrix-extraction prompts and a
    reconciliation table for everything else, so one stub serves both the
    matrix-building and reconciliation calls the endpoint makes.

    For matrix-extraction prompts, echoes back the clause number named in
    MatrixPromptEngine's "ВАЖНО: ... п.N" hint (present for ДС sources, whose
    file name declares the clauses they amend) so ClauseMerger's number-based
    matching behaves realistically; falls back to clause 9 for the general
    policy prompt (no hint — it's the unconstrained baseline extraction).
    """

    def __init__(self):
        self.calls = 0

    def response(self, query: str) -> str:
        self.calls += 1
        if "Пункт (номер/название)" in query:
            hint = re.search(r"п\.(\d+)", query.split("ВАЖНО:", 1)[-1]) if "ВАЖНО:" in query else None
            number = hint.group(1) if hint else "9"
            return (
                "| Пункт (номер/название) | Актуальный текст/значение | Комментарий |\n"
                "|---|---|---|\n"
                f"| {number}. Объект страхования | Оборудование | |\n"
            )
        return (
            "| Наименование поля в декларации | С каким пунктом Ген. полиса сверено | Результат проверки | Комментарий по сверке |\n"
            "|---|---|---|---|\n"
            "| Объект страхования | Объект страхования | совпадает | ОК |\n"
        )


@pytest.fixture
def client(monkeypatch):
    stub = StubModel()
    monkeypatch.setattr(model_mod.ModelFactory, "create", staticmethod(lambda: stub))
    import main
    test_client = TestClient(main.app)
    test_client.stub_model = stub
    return test_client


def _make_policy_folder(tmp_path: Path) -> Path:
    d = Document()
    d.add_paragraph("Полис.")
    d.save(tmp_path / "ГП страхования грузов.docx")

    ds_dir = tmp_path / "ДС"
    ds_dir.mkdir()
    d = Document()
    d.add_paragraph("ДС.")
    d.save(ds_dir / "ДС 1 (п.9).docx")
    return tmp_path


def _make_declaration(folder: Path, number: str) -> Path:
    decl_path = folder / f"{number}.docx"
    d = Document()
    d.add_paragraph(f"Декларация {number}: объект страхования — оборудование.")
    d.save(decl_path)
    return decl_path


class TestReconcileEndpoint:
    def test_full_request_response_cycle(self, tmp_path: Path, client):
        _make_policy_folder(tmp_path)
        decl_path = _make_declaration(tmp_path, "200")

        resp = client.post("/api/reconcile", json={
            "request_id": 1,
            "policy_folder": str(tmp_path),
            "declaration_paths": [str(decl_path)],
        })

        assert resp.status_code == 200
        body = resp.json()
        assert body["matrix"]["clause_count"] == 1
        assert body["matrix"]["cache_hit"] is False

        decl = body["declarations"][0]
        assert decl["declaration_number"] == "200"
        assert decl["type"] == "single"
        assert Path(decl["output_file"]).exists()

    def test_matrix_cache_hit_on_second_request(self, tmp_path: Path, client):
        _make_policy_folder(tmp_path)
        decl_path = _make_declaration(tmp_path, "200")

        # force_rebuild_matrix defaults to True, so the cache is only consulted
        # when the caller explicitly opts out of the rebuild.
        payload = {
            "request_id": 1,
            "policy_folder": str(tmp_path),
            "declaration_paths": [str(decl_path)],
            "force_rebuild_matrix": False,
        }
        client.post("/api/reconcile", json=payload)
        calls_after_first = client.stub_model.calls

        resp = client.post("/api/reconcile", json=payload)

        assert resp.json()["matrix"]["cache_hit"] is True
        # second request should NOT re-extract matrix clauses (only the 1 reconciliation call)
        assert client.stub_model.calls == calls_after_first + 1

    def test_missing_policy_folder_returns_404(self, tmp_path: Path, client):
        resp = client.post("/api/reconcile", json={
            "request_id": 1,
            "policy_folder": str(tmp_path / "missing"),
            "declaration_paths": [],
        })
        assert resp.status_code == 404

    def test_multiple_declarations_each_get_own_output_file(self, tmp_path: Path, client):
        _make_policy_folder(tmp_path)
        paths = [str(_make_declaration(tmp_path, n)) for n in ("200", "201")]

        resp = client.post("/api/reconcile", json={
            "request_id": 1,
            "policy_folder": str(tmp_path),
            "declaration_paths": paths,
        })

        assert resp.status_code == 200
        numbers = {d["declaration_number"] for d in resp.json()["declarations"]}
        assert numbers == {"200", "201"}

    def test_declarations_default_to_folder_when_not_specified(self, tmp_path: Path, client):
        """No declaration_paths at all -> scans {policy_folder}/Декларации/
        recursively (month subfolders included)."""
        _make_policy_folder(tmp_path)
        decl_dir = tmp_path / "Декларации" / "2026-08"
        decl_dir.mkdir(parents=True)
        _make_declaration(decl_dir, "300")

        resp = client.post("/api/reconcile", json={
            "request_id": 1,
            "policy_folder": str(tmp_path),
        })

        assert resp.status_code == 200
        numbers = {d["declaration_number"] for d in resp.json()["declarations"]}
        assert numbers == {"300"}

    def test_declaration_paths_entry_can_be_a_folder(self, tmp_path: Path, client):
        _make_policy_folder(tmp_path)
        decl_dir = tmp_path / "some_folder"
        decl_dir.mkdir()
        _make_declaration(decl_dir, "400")
        _make_declaration(decl_dir, "401")

        resp = client.post("/api/reconcile", json={
            "request_id": 1,
            "policy_folder": str(tmp_path),
            "declaration_paths": [str(decl_dir)],
        })

        assert resp.status_code == 200
        numbers = {d["declaration_number"] for d in resp.json()["declarations"]}
        assert numbers == {"400", "401"}

    def test_no_declarations_found_returns_422(self, tmp_path: Path, client):
        _make_policy_folder(tmp_path)
        resp = client.post("/api/reconcile", json={
            "request_id": 1,
            "policy_folder": str(tmp_path),
        })
        assert resp.status_code == 422

    def test_policy_file_override(self, tmp_path: Path, client):
        """policy_file_override bypasses the "ГП ..." auto-discovery entirely."""
        override_dir = tmp_path / "elsewhere"
        override_dir.mkdir()
        override_path = override_dir / "custom_name.docx"
        d = Document()
        d.add_paragraph("Полис под нестандартным именем.")
        d.save(override_path)

        decl_path = _make_declaration(tmp_path, "200")

        resp = client.post("/api/reconcile", json={
            "request_id": 1,
            "policy_folder": str(tmp_path),  # has no "ГП ..." file at all
            "policy_file_override": str(override_path),
            "declaration_paths": [str(decl_path)],
        })

        assert resp.status_code == 200
        assert resp.json()["matrix"]["clause_count"] == 1

    def test_ds_folder_override(self, tmp_path: Path, client):
        _make_policy_folder(tmp_path)  # already has ДС/ДС 1 (п.9).docx
        custom_ds_dir = tmp_path / "ds_elsewhere"
        custom_ds_dir.mkdir()
        d = Document()
        d.add_paragraph("Другое ДС.")
        d.save(custom_ds_dir / "ДС 5 (п.3).docx")

        decl_path = _make_declaration(tmp_path, "200")

        resp = client.post("/api/reconcile", json={
            "request_id": 1,
            "policy_folder": str(tmp_path),
            "ds_folder_override": str(custom_ds_dir),
            "declaration_paths": [str(decl_path)],
        })

        assert resp.status_code == 200
        # policy clause (9) + the override folder's ДС 5 clause (3) = 2 distinct clauses
        assert resp.json()["matrix"]["clause_count"] == 2
